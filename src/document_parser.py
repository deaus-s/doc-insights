"""
document_parser.py
------------------
Parses PDF, DOCX, TXT, and CSV files into clean text chunks.
Each chunk carries metadata: source filename, page number, chunk index.
"""

import os
import io
import csv
import pandas as pd
from langchain_core.documents import Document


def parse_pdf(file_bytes: bytes, filename: str) -> list[Document]:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={"source": filename, "page": i + 1, "type": "pdf"}
            ))
    return docs


def parse_docx(file_bytes: bytes, filename: str) -> list[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(file_bytes))
    docs = []
    current_section = []
    section_num = 1
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Start new section on headings
        if para.style.name.startswith("Heading") and current_section:
            docs.append(Document(
                page_content="\n".join(current_section),
                metadata={"source": filename, "page": section_num, "type": "docx"}
            ))
            section_num += 1
            current_section = [text]
        else:
            current_section.append(text)
    if current_section:
        docs.append(Document(
            page_content="\n".join(current_section),
            metadata={"source": filename, "page": section_num, "type": "docx"}
        ))
    return docs


def parse_txt(file_bytes: bytes, filename: str) -> list[Document]:
    text = file_bytes.decode("utf-8", errors="ignore")
    # Split into ~1000-char blocks
    chunks = []
    block_size = 1000
    for i in range(0, len(text), block_size):
        chunk = text[i:i + block_size].strip()
        if chunk:
            chunks.append(Document(
                page_content=chunk,
                metadata={"source": filename, "page": i // block_size + 1, "type": "txt"}
            ))
    return chunks


def parse_csv(file_bytes: bytes, filename: str) -> list[Document]:
    text = file_bytes.decode("utf-8", errors="ignore")
    df = pd.read_csv(io.StringIO(text))
    docs = []
    # Convert each row to a readable string
    chunk_size = 20  # rows per chunk
    for i in range(0, len(df), chunk_size):
        chunk_df = df.iloc[i:i + chunk_size]
        content = f"Rows {i+1}–{min(i+chunk_size, len(df))} of {filename}:\n"
        content += chunk_df.to_string(index=False)
        docs.append(Document(
            page_content=content,
            metadata={"source": filename, "page": i // chunk_size + 1, "type": "csv"}
        ))
    return docs


def parse_document(file_bytes: bytes, filename: str) -> list[Document]:
    """Main entry point — routes to correct parser based on extension."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes, filename)
    elif ext == ".docx":
        return parse_docx(file_bytes, filename)
    elif ext == ".txt":
        return parse_txt(file_bytes, filename)
    elif ext == ".csv":
        return parse_csv(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
